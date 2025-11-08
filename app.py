# app.py — Mini-Grammarly (Streamlit) with Plagiarism Checker + Modern UI
# Run locally:  streamlit run app.py

import difflib
import io
import os
import json
import time
import textwrap
import warnings
from dataclasses import dataclass
from typing import List, Tuple, Optional

import streamlit as st
import requests

# -----------------------------------------------------------------------------
# Page + housekeeping
# -----------------------------------------------------------------------------
st.set_page_config(page_title="Mini-Grammarly", page_icon="📝")
warnings.filterwarnings("ignore", message="pkg_resources is deprecated")

# ---- CSS injection (Bootstrap-like look) ----
def _load_css():
    path = os.path.join(os.path.dirname(__file__), "assets", "styles.css")
    try:
        with open(path, "r", encoding="utf-8") as f:
            st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)
    except Exception:
        pass

_load_css()

# Optional libs for readability and exports
try:
    import textstat
except Exception:
    textstat = None

try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import docx  # python-docx
    from docx.shared import Pt
    from docx.enum.text import WD_COLOR_INDEX
except Exception:
    docx = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.pdfgen import canvas
    _pdf_ok = True
except Exception:
    _pdf_ok = False

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    _sk_ok = True
except Exception:
    _sk_ok = False

# language_tool_python (optional for Local server mode)
try:
    import language_tool_python as lt
    lt_error = None
except Exception as e:
    lt = None
    lt_error = repr(e)

# Small exception for public API rate limit
class APIRateLimit(Exception):
    pass

def _rep_value(x):
    """Return replacement string from either object w/.value or dict {'value':...}."""
    if hasattr(x, "value"):
        return x.value
    if isinstance(x, dict):
        return x.get("value", "")
    return str(x or "")

# -----------------------------------------------------------------------------
# Secrets helper (works with/without .streamlit/secrets.toml)
# -----------------------------------------------------------------------------
def _get_secret(name: str, default: str = "") -> str:
    try:
        return st.secrets.get(name, default)
    except Exception:
        return os.environ.get(name, default)

REMOTE = _get_secret("LT_REMOTE_SERVER_URL", "")
HF_TOKEN = _get_secret("HF_API_TOKEN", "")
PLAG_API_KEY = _get_secret("PLAG_API_KEY", "")

# -----------------------------------------------------------------------------
# Paraphrasing (API first; optional local)
# -----------------------------------------------------------------------------
def paraphrase_api(text: str, tone: str = "", model: str = "Vamsi/T5_Paraphrase_Paws") -> str:
    token = HF_TOKEN
    if not token:
        return "(Paraphraser via API unavailable — set HF_API_TOKEN in secrets.)"

    url = f"https://api-inference.huggingface.co/models/{model}"
    headers = {"Authorization": f"Bearer {token}"}
    tone_map = {"Formal": "Rewrite formally:", "Friendly": "Rewrite in a friendly tone:", "Concise": "Rewrite concisely:", "Neutral": ""}
    style = tone_map.get(tone or "Neutral", "")
    payload = {"inputs": f"paraphrase: {style} {text}".strip(), "options": {"wait_for_model": True}}

    for _ in range(3):
        r = requests.post(url, headers=headers, data=json.dumps(payload), timeout=60)
        if r.status_code == 503:
            time.sleep(2)
            continue
        r.raise_for_status()
        out = r.json()
        if isinstance(out, list) and out:
            if isinstance(out[0], dict) and "generated_text" in out[0]:
                return out[0]["generated_text"]
            if isinstance(out[0], dict) and "summary_text" in out[0]:
                return out[0]["summary_text"]
        return str(out)
    return "(Paraphrase API is warming up. Try again.)"

_paraphraser = None
_paraphraser_tokenizer = None
_PARA_MODEL = "Vamsi/T5_Paraphrase_Paws"

def _ensure_paraphraser() -> bool:
    global _paraphraser, _paraphraser_tokenizer
    if _paraphraser is not None and _paraphraser_tokenizer is not None:
        return True
    try:
        from transformers import AutoModelForSeq2SeqLM, AutoTokenizer
        _paraphraser_tokenizer = AutoTokenizer.from_pretrained(_PARA_MODEL)
        _paraphraser = AutoModelForSeq2SeqLM.from_pretrained(_PARA_MODEL)
        return True
    except Exception:
        return False

def paraphrase_local(text: str, beams: int = 5, max_len: int = 256, style_hint: str = "") -> str:
    if not _ensure_paraphraser():
        return "(Paraphraser unavailable — install transformers, sentencepiece, and torch.)"
    prefix = "paraphrase: " + (style_hint + " " if style_hint else "") + text
    inputs = _paraphraser_tokenizer.encode(prefix, return_tensors="pt", truncation=True)
    outputs = _paraphraser.generate(
        inputs, max_length=max_len, num_beams=beams, early_stopping=True, no_repeat_ngram_size=3
    )
    return _paraphraser_tokenizer.decode(outputs[0], skip_special_tokens=True)

# -----------------------------------------------------------------------------
# UI
# -----------------------------------------------------------------------------
st.title("📝 Mini-Grammarly")
st.caption("Grammar, style, rewrite & plagiarism — with a clean, modern look")

SUPPORTED_LANGS = {
    "English (US)": "en-US",
    "English (UK)": "en-GB",
    "German": "de-DE",
    "French": "fr-FR",
    "Spanish": "es",
    "Portuguese": "pt-PT",
}

with st.sidebar:
    st.header("Settings")
    st.caption(
        f"LangTool pkg: {'✅' if lt else '❌'}" + (f" (import error: {lt_error})" if lt_error else "")
    )
    st.caption(f"HF API token: {'✅' if HF_TOKEN else '❌'}")
    st.caption("Remote LT server: " + ("✅ (from secret/env)" if REMOTE else "❌ (using Public/Local)"))

    lang = st.selectbox("Language", list(SUPPORTED_LANGS.keys()), index=0)
    lang_code = SUPPORTED_LANGS[lang]

    ENGINE = st.radio(
        "Grammar engine",
        ["Public API (free, rate-limited)", "Local server (Java)", "Remote server (from secret/env)"],
        index=0 if not REMOTE else 2,
        help="Public is easiest; Local needs Java; Remote uses your own LanguageTool server URL.",
    )

    do_rewrite = st.checkbox("Enable paraphrase/rewrite", value=False)
    target_tone = st.selectbox("Target tone (rewrite)", ["Neutral", "Formal", "Friendly", "Concise"], index=0)
    beams = st.slider("Paraphrase quality (beams, local only)", 2, 8, 5)

    st.divider()
    st.subheader("Plagiarism checker")
    enable_plag = st.checkbox("Enable plagiarism checking", value=False)
    plag_mode = st.radio(
        "Mode",
        ["Local (upload refs / paste text)", "External API (bring your own key)"],
        index=0,
        help="Local: compare against your uploads and pasted references; API: wire a provider key.",
    )
    plag_thresh = st.slider("Flag similarity ≥", 0.0, 1.0, 0.35, 0.01)
    pasted_refs = st.text_area(
        "Paste reference text(s) (optional). Separate multiple sources with a line containing only `---`",
        height=120,
        placeholder="Paste any text that you want to compare against.\n---\nPaste another reference...",
    )
    uploaded_files = st.file_uploader(
        "Upload reference files (TXT, PDF, DOCX) — multiple allowed",
        type=["txt", "pdf", "docx"],
        accept_multiple_files=True,
    )

text = st.text_area("Paste or type your text", height=220, placeholder="Type here…")

c1, c2, c3, c4 = st.columns([1, 1, 1, 1])
with c1:
    run = st.button("Check grammar & spelling")
with c2:
    rewrite_btn = st.button("Paraphrase / Rewrite")
with c3:
    stats_btn = st.button("Readability report")
with c4:
    plag_btn = st.button("Run plagiarism check")

# -----------------------------------------------------------------------------
# Grammar checking
# -----------------------------------------------------------------------------
@dataclass
class Issue:
    rule: str
    message: str
    offset: int
    length: int
    context: str
    replacements: List[str]

def _lt_check_http(text: str, lang_code: str, remote_url: Optional[str]) -> Tuple[str, List[Issue]]:
    """LanguageTool /v2/check via HTTP (public or remote) with robust fallbacks."""
    if not isinstance(text, str) or not text.strip():
        return text, []

    # keep text sane for public endpoint
    MAX_CHARS = 20000
    payload_text = text[:MAX_CHARS]

    base = (remote_url.rstrip("/") if remote_url else "https://api.languagetool.org") + "/v2/check"
    headers = {
        "User-Agent": "mini-grammarly/1.1 (streamlit)",
        "Accept": "application/json",
    }

    def _try(language_value: str):
        data = {
            "text": payload_text,
            "language": language_value,
            "enabledOnly": "true",
        }
        r = requests.post(base, data=data, headers=headers, timeout=30)
        if r.status_code == 429:
            raise APIRateLimit("Public API rate limit hit")
        if r.status_code >= 400:
            try:
                msg = r.json()
            except Exception:
                msg = r.text
            raise requests.HTTPError(f"{r.status_code} {r.reason}: {msg}", response=r)
        try:
            return r.json()
        except Exception:
            raise requests.HTTPError(f"Unexpected response from LanguageTool: {r.text[:200]}", response=r)

    # try chosen language; on 400 retry fallbacks
    try_order = [lang_code]
    if lang_code not in ("auto", "auto-en"):
        try_order += ["auto", "auto-en", "en-US"]

    last_err = None
    for lang_try in try_order:
        try:
            js = _try(lang_try)
            matches = js.get("matches", [])
            issues: List[Issue] = []
            corrected = list(text)
            for m in sorted(matches, key=lambda x: x.get("offset", 0), reverse=True):
                offset = m.get("offset", 0)
                length = m.get("length", 0)
                message = m.get("message", "")
                rule_id = (m.get("rule") or {}).get("id", "Rule")
                context = (m.get("context") or {}).get("text", "")
                repls = [rep.get("value") for rep in m.get("replacements", []) if "value" in rep]
                issues.append(Issue(rule=rule_id, message=message, offset=offset, length=length, context=context, replacements=repls))
                if repls:
                    suggestion = repls[0]
                    corrected[offset:offset+length] = list(suggestion)
            corrected_text = "".join(corrected)
            return corrected_text, list(reversed(issues))
        except requests.HTTPError as e:
            last_err = e
            if e.response is None or e.response.status_code not in (400,):
                raise

    raise last_err if last_err else requests.HTTPError("LanguageTool request failed with unknown error")

def get_tool(lang_code: str):
    if lt is None:
        return None
    try:
        if ENGINE.startswith("Remote"):
            if not REMOTE:
                return None
            return lt.LanguageTool(lang_code, remote_server=REMOTE)
        if ENGINE.startswith("Public"):
            return None  # we force HTTP path for Public
        return lt.LanguageTool(lang_code)  # Local (Java)
    except Exception:
        return None

def check_text(text: str, lang_code: str) -> Tuple[str, List[Issue]]:
    """
    Grammar check with three modes:
      - Public API: direct HTTP (avoids 426/JSON issues)
      - Remote server: HTTP to your LT server (LT_REMOTE_SERVER_URL)
      - Local server: try language_tool_python (Java), else HTTP
    """
    if ENGINE.startswith("Public"):
        return _lt_check_http(text, lang_code, remote_url=None)

    if ENGINE.startswith("Remote"):
        return _lt_check_http(text, lang_code, remote_url=REMOTE if REMOTE else None)

    if lt is not None:
        tool = get_tool(lang_code)
        if tool is not None:
            try:
                matches = tool.check(text)
            except Exception:
                return _lt_check_http(text, lang_code, remote_url=None)

            issues: List[Issue] = []
            for m in matches:
                rep_list = getattr(m, "replacements", []) or []
                repl = [_rep_value(r) for r in rep_list]
                issues.append(
                    Issue(
                        rule=getattr(getattr(m, "rule", None), "id", "Rule"),
                        message=getattr(m, "message", ""),
                        offset=getattr(m, "offset", 0),
                        length=getattr(m, "errorLength", 0),
                        context=getattr(m, "context", ""),
                        replacements=repl,
                    )
                )

            corrected = list(text)
            for m in sorted(matches, key=lambda x: getattr(x, "offset", 0), reverse=True):
                rep_list = getattr(m, "replacements", []) or []
                if rep_list:
                    suggestion = _rep_value(rep_list[0])
                    start = getattr(m, "offset", 0)
                    end = start + getattr(m, "errorLength", 0)
                    corrected[start:end] = list(suggestion)

            corrected_text = "".join(corrected)
            return corrected_text, issues

    return _lt_check_http(text, lang_code, remote_url=None)

@st.cache_data(show_spinner=False, ttl=300)
def check_text_cached(text: str, lang_code: str, engine_key: str, remote_key: str):
    return check_text(text, lang_code)

def highlight_diff(before: str, after: str) -> str:
    diff = difflib.ndiff(before.split(), after.split())
    html_tokens = []
    for token in diff:
        if token.startswith("- "):
            html_tokens.append(f"<del>{token[2:]}</del>")
        elif token.startswith("+ "):
            html_tokens.append(f"<ins>{token[2:]}</ins>")
        elif token.startswith("? "):
            continue
        else:
            html_tokens.append(token[2:])
    return " ".join(html_tokens)

# -----------------------------------------------------------------------------
# Plagiarism (local)
# -----------------------------------------------------------------------------
def _read_txt(bytes_data: bytes) -> str:
    try:
        return bytes_data.decode("utf-8")
    except Exception:
        try:
            return bytes_data.decode("latin-1")
        except Exception:
            return ""

def _read_pdf(file) -> str:
    if pdfplumber is None:
        return ""
    try:
        text = []
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text.append(page.extract_text() or "")
        return "\n".join(text)
    except Exception:
        return ""

def _read_docx(file) -> str:
    if docx is None:
        return ""
    try:
        d = docx.Document(file)
        return "\n".join(p.text for p in d.paragraphs)
    except Exception:
        return ""

def _extract_text_from_upload(upload) -> str:
    name = upload.name.lower()
    data = upload.read()
    upload.seek(0)
    if name.endswith(".txt"):
        return _read_txt(data)
    if name.endswith(".pdf"):
        return _read_pdf(io.BytesIO(data))
    if name.endswith(".docx"):
        return _read_docx(io.BytesIO(data))
    return ""

def _split_pasted_refs(pasted: str) -> List[str]:
    if not pasted.strip():
        return []
    parts, current = [], []
    for line in pasted.splitlines():
        if line.strip() == "---":
            if current:
                parts.append("\n".join(current).strip())
                current = []
        else:
            current.append(line)
    if current:
        parts.append("\n".join(current).strip())
    return [p for p in parts if p]

def _cosine_similarity(a: str, b: str) -> float:
    if not _sk_ok:
        return 0.0
    try:
        vec = TfidfVectorizer(analyzer="char", ngram_range=(3, 5), min_df=1)
        X = vec.fit_transform([a, b])
        return float(cosine_similarity(X[0], X[1])[0][0])
    except Exception:
        return 0.0

def _highlight_overlap(a: str, b: str, min_match: int = 30) -> Tuple[str, List[str]]:
    sm = difflib.SequenceMatcher(None, a, b)
    matches = [m for m in sm.get_matching_blocks() if m.size >= min_match]
    merged = []
    for m in matches:
        if merged and m.a <= merged[-1].a + merged[-1].size + 10:
            last = merged[-1]
            merged[-1] = difflib.Match(last.a, last.b, (m.a + m.size) - last.a)
        else:
            merged.append(m)
    out, last = [], 0
    matched_b = []
    for m in merged:
        out.append(a[last:m.a])
        out.append(f"<mark>{a[m.a:m.a+m.size]}</mark>")
        matched_b.append(b[m.b:m.b+m.size])
        last = m.a + m.size
    out.append(a[last:])
    return "".join(out), matched_b

def run_local_plagiarism(target: str, ref_texts: List[Tuple[str, str]], thresh: float = 0.35):
    results = []
    for name, src in ref_texts:
        if not src.strip():
            continue
        sim = _cosine_similarity(target, src)
        if sim >= thresh:
            highlighted, overlaps = _highlight_overlap(target, src, min_match=30)
            results.append({
                "Source": name,
                "Similarity": round(sim, 3),
                "OverlapsFound": len(overlaps),
                "HighlightedTargetHTML": highlighted,
                "OverlapSamples": overlaps[:5]
            })
    results.sort(key=lambda x: x["Similarity"], reverse=True)
    return results

# Exports
def _build_docx_report(target_text: str, results: List[dict]) -> bytes:
    if docx is None:
        return b""
    d = docx.Document()
    d.add_heading("Plagiarism Report", level=1)
    p = d.add_paragraph("Summary: ")
    run = p.add_run(f"{len(results)} match(es) above threshold.")
    run.font.bold = True

    d.add_heading("Checked text", level=2)
    d.add_paragraph(target_text)

    for r in results:
        d.add_heading(f"Source: {r['Source']}  (similarity {r['Similarity']})", level=2)
        d.add_paragraph(f"Overlaps found: {r['OverlapsFound']}")
        para = d.add_paragraph()
        tmp = r["HighlightedTargetHTML"].replace("<mark>", "\u0001").replace("</mark>", "\u0002")
        pieces, buff, marked = [], "", False
        for ch in tmp:
            if ch == "\u0001":
                if buff:
                    pieces.append((buff, False)); buff = ""
                marked = True
            elif ch == "\u0002":
                if buff:
                    pieces.append((buff, True)); buff = ""
                marked = False
            else:
                buff += ch
        if buff:
            pieces.append((buff, marked))
        for text_part, is_marked in pieces:
            run = para.add_run(text_part)
            run.font.size = Pt(11)
            if is_marked:
                try:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                except Exception:
                    pass

        if r["OverlapSamples"]:
            d.add_paragraph("Sample matching fragments from source:")
            for frag in r["OverlapSamples"]:
                d.add_paragraph(f"• {frag}")

    bio = io.BytesIO()
    d.save(bio)
    return bio.getvalue()

def _build_pdf_report(target_text: str, results: List[dict]) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.pdfgen import canvas
    except Exception:
        return b""
    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=A4)
    width, height = A4
    x, y = 2*cm, height - 2*cm
    line_h = 12

    def draw_line(text, bold=False):
        nonlocal y
        if y < 2*cm:
            c.showPage()
            y = height - 2*cm
        c.setFont("Helvetica-Bold" if bold else "Helvetica", 10)
        c.drawString(x, y, text)
        y -= line_h

    draw_line("Plagiarism Report", bold=True)
    draw_line(f"Matches above threshold: {len(results)}")
    y -= 6
    draw_line("Checked text:", bold=True)
    for ln in textwrap.wrap(target_text, width=95):
        draw_line(ln)
    y -= 6
    for r in results:
        draw_line(f"Source: {r['Source']}  (similarity {r['Similarity']})", bold=True)
        draw_line(f"Overlaps found: {r['OverlapsFound']}")
        if r["OverlapSamples"]:
            draw_line("Samples:")
            for frag in r["OverlapSamples"]:
                for ln in textwrap.wrap(f"• {frag}", width=95):
                    draw_line(ln)
        y -= 6

    c.showPage()
    c.save()
    return bio.getvalue()

def run_external_plagiarism_stub(target: str, api_key: str) -> List[dict]:
    return []

# -----------------------------------------------------------------------------
# Actions
# -----------------------------------------------------------------------------
if run and text.strip():
    try:
        fixed, issues = check_text_cached(text, lang_code, ENGINE, REMOTE)
    except APIRateLimit:
        st.error(
            "You’ve hit the free LanguageTool API rate limit.\n\n"
            "• Try again later, or\n"
            "• Set LT_REMOTE_SERVER_URL for your own server, or\n"
            "• Use Local server on your PC (Java required)."
        )
        st.stop()
    except Exception:
        import traceback
        st.error("Grammar check failed:\n\n" + traceback.format_exc())
        st.stop()

    if not issues:
        st.success("No issues found.")

    st.subheader("Suggested fixes (auto-applied preview)")
    st.markdown(highlight_diff(text, fixed), unsafe_allow_html=True)

    if issues:
        rows = []
        for i, it in enumerate(issues, 1):
            bad = text[it.offset : it.offset + it.length]
            top_suggestion = it.replacements[0] if it.replacements else "(none)"
            rows.append({"#": i, "Rule": it.rule, "Message": it.message, "Text": bad, "Suggestion": top_suggestion,
                         "Pos": f"{it.offset}:{it.offset + it.length}"})
        try:
            import pandas as pd
            st.dataframe(pd.DataFrame(rows), use_container_width=True)
        except Exception:
            for r in rows:
                st.write(r)

    st.download_button("Download corrected text", fixed.encode("utf-8"), file_name="corrected.txt")

if rewrite_btn and text.strip():
    if not do_rewrite:
        st.subheader("Paraphrased / Rewritten")
        st.write("(Rewrite disabled in sidebar.)")
    else:
        out = paraphrase_api(text, tone=target_tone)
        if out.startswith("("):  # API not configured; fallback to local if available
            style_hint = {"Formal":"Rewrite formally:", "Friendly":"Rewrite in a friendly tone:", "Concise":"Rewrite concisely:"}.get(target_tone, "")
            out = paraphrase_local(text, beams=beams, style_hint=style_hint)
        st.subheader("Paraphrased / Rewritten")
        st.write(out)
        if "unavailable" in out:
            st.info("To enable local paraphrasing, add to requirements.txt: transformers, sentencepiece, torch")

if stats_btn and text.strip():
    st.subheader("Readability")
    if textstat is None:
        st.info("Install textstat for readability scores: pip install textstat")
    else:
        cA, cB = st.columns(2)
        with cA:
            st.metric("Flesch Reading Ease", f"{textstat.flesch_reading_ease(text):.1f}")
            st.metric("Flesch–Kincaid Grade", f"{textstat.flesch_kincaid_grade(text):.1f}")
            st.metric("Gunning Fog", f"{textstat.gunning_fog(text):.1f}")
        with cB:
            st.metric("SMOG Index", f"{textstat.smog_index(text):.1f}")
            st.metric("Coleman–Liau", f"{textstat.coleman_liau_index(text):.1f}")
            st.metric("Dale–Chall", f"{textstat.dale_chall_readability_score(text):.1f}")

if plag_btn and enable_plag and text.strip():
    st.subheader("Plagiarism results")

    results: List[dict] = []
    if plag_mode.startswith("Local"):
        if not _sk_ok:
            st.warning("Local plagiarism needs scikit-learn. Add to requirements.txt: scikit-learn")
            st.stop()
        references: List[Tuple[str, str]] = []
        for up in uploaded_files or []:
            content = _extract_text_from_upload(up)
            references.append((up.name, content))
        for i, block in enumerate(_split_pasted_refs(pasted_refs), 1):
            references.append((f"Pasted reference {i}", block))

        if not references:
            st.info("Provide at least one reference (upload files and/or paste text) to compare against.")
        else:
            results = run_local_plagiarism(text, references, thresh=plag_thresh)

    else:
        if not PLAG_API_KEY:
            st.info("Set PLAG_API_KEY in secrets to enable External API mode.")
        else:
            results = run_external_plagiarism_stub(text, PLAG_API_KEY)

    if not results:
        st.success("No suspicious similarity above the threshold.")
    else:
        try:
            import pandas as pd
            df = pd.DataFrame([{"Source": r["Source"], "Similarity": r["Similarity"], "Overlaps": r["OverlapsFound"]} for r in results])
            st.dataframe(df, use_container_width=True)
            csv = df.to_csv(index=False).encode("utf-8")
            st.download_button("Download CSV report", csv, "plagiarism_report.csv", "text/csv")
        except Exception:
            for r in results:
                st.write({"Source": r["Source"], "Similarity": r["Similarity"], "Overlaps": r["OverlapsFound"]})

        for r in results:
            with st.expander(f"Details — {r['Source']} (similarity {r['Similarity']})"):
                st.markdown(r["HighlightedTargetHTML"], unsafe_allow_html=True)
                if r["OverlapSamples"]:
                    st.caption("Sample matching fragments from source:")
                    for frag in r["OverlapSamples"]:
                        st.write(f"• {frag}")

        colA, colB = st.columns(2)
        with colA:
            docx_bytes = _build_docx_report(text, results)
            st.download_button("Download DOCX report", docx_bytes, "plagiarism_report.docx",
                               "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               disabled=(not docx_bytes))
        with colB:
            pdf_bytes = _build_pdf_report(text, results)
            st.download_button("Download PDF report", pdf_bytes, "plagiarism_report.pdf", "application/pdf",
                               disabled=(not pdf_bytes))

st.markdown(
    '<div class="app-footer">Made with Streamlit · LanguageTool · Hugging Face · scikit-learn</div>',
    unsafe_allow_html=True
)

st.caption(
    "Grammar engines: Public API (rate-limited), Local (Java), Remote (set LT_REMOTE_SERVER_URL). "
    "Paraphrasing: uses Hugging Face Inference API via HF_API_TOKEN; falls back to local Transformers if installed. "
    "Plagiarism (Local): compares against uploads/pasted refs with TF-IDF cosine similarity and overlap highlights."
)
